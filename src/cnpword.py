from docx import Document
from PyQt5.QtCore import QDate
import cnpdb, cnputil, cnpval
    
def save_client_to_word(clients, filepath):
    doc = Document()
    doc.add_heading("Client Information", level=1)

    # --- Names section ---
    doc.add_heading("Names", level=2)
    table = doc.add_table(rows=3, cols=4)
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0].cells
    hdr[0].text = "Type"
    hdr[1].text = "First"
    hdr[2].text = "Middle"
    hdr[3].text = "Last"

    # Korean
    row = table.rows[1].cells
    row[0].text = "Korean"
    row[1].text = cnputil.safe_val(clients.get(cnpdb.col_first_name_kor, ""))
    row[3].text = cnputil.safe_val(clients.get(cnpdb.col_last_name_kor, ""))

    # English
    row = table.rows[2].cells
    row[0].text = "English"
    row[1].text = cnputil.safe_val(clients.get(cnpdb.col_first_name_eng, ""))
    row[2].text = cnputil.safe_val(clients.get(cnpdb.col_middle_name_eng, ""))
    row[3].text = cnputil.safe_val(clients.get(cnpdb.col_last_name_eng, ""))

    # --- DOB/AGE/SEX row ---
    doc.add_heading("Personal Info", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    row = table.rows[0].cells
    row[0].text = f"DOB: {cnputil.safe_val(clients.get(cnpdb.col_dob,''))}"
    qdate_obj = QDate.fromString(clients.get(cnpdb.col_dob,''), cnpval.date_format_r)
    cur_age = cnputil.calculate_age(qdate_obj)
    row[1].text = f"Age: {cur_age}"
    row[2].text = f"Sex: {cnputil.safe_val(clients.get(cnpdb.col_sex,''))}"

    # --- Assessments ---
    doc.add_heading("Assessments", level=2)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    labels = ["Initial Assessment", "14th Assessment", "90th Assessment", "Change Assessment"]
    keys   = [cnpdb.col_initial_assessment, cnpdb.col_assessment_14th, 
              cnpdb.col_assessment_90th, cnpdb.col_change_assessment]

    for i, (lbl, key) in enumerate(zip(labels, keys)):
        row = table.rows[i].cells
        row[0].text = lbl
        row[1].text = cnputil.safe_val(clients.get(key, ""))

    # Done flag
    change_assement_done = cnputil.safe_val(clients.get(cnpdb.col_change_assessment_done))
    table.add_row().cells[0].text = "Done"
    if change_assement_done == cnpdb.change_assessment_yes:
        table.rows[-1].cells[1].text = cnpdb.change_assessment_yes
    else:
        table.rows[-1].cells[1].text = cnpdb.change_assessment_no
        
    # --- Room + Comments ---
    doc.add_heading("Other Info", level=2)
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    row = table.rows[0].cells
    row[0].text = "Room Number"
    row[1].text = cnputil.safe_val(clients.get(cnpdb.col_room_number, ""))
    row = table.rows[1].cells
    row[0].text = "Comments"
    row[1].text = cnputil.safe_val(clients.get(cnpdb.col_comments, ""))

    doc.save(filepath)
